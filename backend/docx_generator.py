from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import re
from datetime import datetime

def find_page_for_text(search_text, paginated_text):
    if not search_text or search_text == "Not Available":
        return None
    
    clean_search = search_text.upper()
    pag_upper = paginated_text.upper()
    
    # 1. Direct match
    idx = pag_upper.find(clean_search)
    
    # 2. Key word match (first 2 words)
    if idx == -1:
        words = [w for w in re.split(r'\W+', clean_search) if len(w) > 3]
        if words:
            idx = pag_upper.find(words[0])
    
    if idx == -1:
        return None
    
    # find all page markers
    matches = list(re.finditer(r"--- \[START OF PAGE (\d+)\] ---", paginated_text))
    last_page = None
    for m in matches:
        if m.start() < idx:
            last_page = int(m.group(1))
        else:
            break
    return last_page

def add_table_with_headers(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

def generate_report_docx(json_data, inspection_text, thermal_text, inspection_images, thermal_images):
    doc = Document()
    
    # 1. Cover page
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Detailed Diagnostic Report')
    run.font.size = Pt(24)
    run.font.bold = True
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(f'Date: {datetime.now().strftime("%Y-%m-%d")}')
    date_run.font.size = Pt(14)
    doc.add_page_break()
    
    # 2. Property Issue Summary
    doc.add_heading('Property Issue Summary', level=1)
    doc.add_paragraph(json_data.get("property_issue_summary", "Not Available"))
    
    # 3. Area-wise Observations
    doc.add_heading('Area-wise Observations', level=1)
    for obs in json_data.get("area_wise_observations", []):
        area = obs.get("area", "Unknown Area")
        doc.add_heading(area, level=2)
        doc.add_paragraph(f"Observation: {obs.get('observation', 'Not Available')}")
        
        thermal_finding = obs.get("thermal_finding", "Not Available")
        if thermal_finding != "Not Available":
            doc.add_paragraph(f"Thermal Finding: {thermal_finding}")
            
        found_image = False
        
        # Inspection image
        insp_page = find_page_for_text(area, inspection_text)
        if insp_page and insp_page in inspection_images and len(inspection_images[insp_page]) > 0:
            img = inspection_images[insp_page].pop(0)
            try:
                doc.add_picture(io.BytesIO(img["bytes"]), width=Inches(4.0))
                found_image = True
            except Exception:
                pass
                
        # Thermal image
        therm_page = find_page_for_text(area, thermal_text)
        if therm_page and therm_page in thermal_images and len(thermal_images[therm_page]) > 0:
            img = thermal_images[therm_page].pop(0)
            try:
                doc.add_picture(io.BytesIO(img["bytes"]), width=Inches(4.0))
                found_image = True
            except Exception:
                pass
                
        if not found_image:
            doc.add_paragraph("[Image Not Available]")
            
    # 4. Probable Root Cause
    doc.add_heading('Probable Root Cause', level=1)
    rc_list = json_data.get("probable_root_cause", [])
    if rc_list and isinstance(rc_list, list):
        add_table_with_headers(doc, ["Issue", "Root Cause"], [[rc.get("issue", ""), rc.get("root_cause", "")] for rc in rc_list])
    else:
        doc.add_paragraph("Not Available")
        
    # 5. Severity Assessment
    doc.add_heading('Severity Assessment', level=1)
    sev_list = json_data.get("severity_assessment", [])
    if sev_list and isinstance(sev_list, list):
        add_table_with_headers(doc, ["Issue", "Severity", "Reasoning"], [[s.get("issue", ""), s.get("severity", ""), s.get("reasoning", "")] for s in sev_list])
    else:
        doc.add_paragraph("Not Available")
        
    # 6. Recommended Actions
    doc.add_heading('Recommended Actions', level=1)
    act_list = json_data.get("recommended_actions", [])
    if act_list and isinstance(act_list, list):
        add_table_with_headers(doc, ["Issue", "Action", "Priority"], [[a.get("issue", ""), a.get("action", ""), a.get("priority", "")] for a in act_list])
    else:
        doc.add_paragraph("Not Available")
        
    # 7. Additional Notes
    doc.add_heading('Additional Notes', level=1)
    notes = json_data.get("additional_notes", "Not Available")
    if not notes:
        notes = "Not Available"
    doc.add_paragraph(str(notes))
    
    # 8. Missing or Unclear Information
    doc.add_heading('Missing or Unclear Information', level=1)
    missing = json_data.get("missing_or_unclear_information", [])
    if isinstance(missing, list) and missing:
        for item in missing:
            doc.add_paragraph(str(item), style='List Bullet')
    else:
        doc.add_paragraph(str(missing))
        
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream
